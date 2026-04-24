"""
Generate a DOCX resume from a Word template and a markdown source file.
Run from the repo root: python scripts/build_docx_from_template.py [template] [source] [output]
  e.g.  python scripts/build_docx_from_template.py "templates/Excella Resume Template 2026.docx" resume.md resume_from_template.docx

Markdown → Word style mapping
------------------------------
  # Name                 → Name  (fallback: Title → Heading 1 → Normal)
  ## Section             → Section Header  (fallback: Heading 1 → Normal)
  ### Company / Role     → Company  (fallback: Heading 2 → Normal)
  #### Sub-role          → Normal italic
  **bold** (pre-section) → Job Position or Title  (fallback: Subtitle → Normal)
  **bold** (in body)     → Normal italic  (date lines etc.)
  **partial bold**       → Normal
  *italic*               → Normal italic
  - bullet               → Bullet Points  (fallback: List Paragraph → Normal)
  | table row |          → Skills  (fallback: Normal)
  plain text             → Normal

Style resolution
----------------
All named styles are resolved with fallbacks so the script degrades gracefully
when a template uses different style names. Run generate_md_from_template.py to
inspect the available styles in any template revision.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

TEMPLATE = sys.argv[1] if len(sys.argv) > 1 else "templates/Excella Resume Template 2026.docx"
SOURCE   = sys.argv[2] if len(sys.argv) > 2 else "resume.md"
OUTPUT   = sys.argv[3] if len(sys.argv) > 3 else "resume_from_template.docx"

for f in [TEMPLATE, SOURCE]:
    if not Path(f).exists():
        sys.exit(f"ERROR: file not found: {f}")

tmpl = Document(TEMPLATE)
sect_pr = deepcopy(tmpl.element.body.find(qn("w:sectPr")))

doc = Document(TEMPLATE)
body = doc.element.body
for child in list(body):
    body.remove(child)

# ── Style resolution with fallbacks ─────────────────────────────────────────
_available = {s.name for s in doc.styles}
_FALLBACKS = {
    "Name":                  ["Name", "Title", "Heading 1", "Normal"],
    "Section Header":        ["Section Header", "Heading 1", "Normal"],
    "Company":               ["Company", "Heading 2", "Normal"],
    "Job Position or Title": ["Job Position or Title", "Subtitle", "Normal"],
    "Bullet Points":         ["Bullet Points", "List Paragraph", "Normal"],
    "Skills":                ["Skills", "Normal"],
    "Certifications":        ["Certifications", "Normal"],
    "Section Subheader":     ["Section Subheader", "Heading 3", "Normal"],
}


def resolve(style_name):
    for candidate in _FALLBACKS.get(style_name, [style_name, "Normal"]):
        if candidate in _available:
            return candidate
    return "Normal"


def ap(text, style, italic=False):
    para = doc.add_paragraph(style=resolve(style))
    if text:
        run = para.add_run(text)
        if italic:
            run.italic = True
    return para


def strip_bold(text):
    return re.sub(r"\*\*([^*]*)\*\*", r"\1", text)


def strip_links(text):
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)


def clean(text):
    return strip_links(strip_bold(text))


lines = Path(SOURCE).read_text(encoding="utf-8").splitlines()

in_header = True   # True until the first ## section; controls bold-line style

for line in lines:
    s = line.strip()

    if not s or s == "---":
        continue

    # H1 → Name
    if s.startswith("# "):
        ap(s[2:], "Name")
        continue

    # H2 → Section Header
    if s.startswith("## "):
        ap(s[3:], "Section Header")
        in_header = False
        continue

    # H3 → Company
    if s.startswith("### "):
        ap(clean(s[4:]), "Company")
        continue

    # H4 → italic sub-role / date range
    if s.startswith("#### "):
        ap(clean(s[5:]), "Normal", italic=True)
        continue

    # Fully bold line: **text**
    if s.startswith("**") and s.endswith("**"):
        text = clean(s)
        if in_header:
            ap(text, "Job Position or Title")
        else:
            ap(text, "Normal", italic=True)
        continue

    # Partially bold line: **Label:** value
    if s.startswith("**") and "**" in s[2:]:
        ap(clean(s), "Normal")
        continue

    # Italic line: *text*
    if re.match(r"^\*[^*]", s) and s.endswith("*"):
        ap(s[1:-1], "Normal", italic=True)
        continue

    # Table separator: skip
    if s.startswith("|") and not s.replace("|", "").replace("-", "").replace(" ", ""):
        continue

    # Table data row → Skills  (only structure that produces tables in the skeleton)
    if s.startswith("|"):
        cols = [clean(c.strip()) for c in s.split("|") if c.strip()]
        text = ": ".join(cols)
        if text:
            ap(text, "Skills")
        continue

    # Bullet item
    if s.startswith("- "):
        ap(clean(s[2:]), "Bullet Points")
        continue

    # Plain text → Normal
    ap(clean(s), "Normal")

body.append(sect_pr)

doc.save(OUTPUT)
print(f"Built: {OUTPUT}")

